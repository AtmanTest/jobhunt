from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.2)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

def add_title(text, size=22, color=RGBColor(0x11, 0x11, 0x11), bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    return p

def add_mixed_line(parts):
    """parts = [(text, size, color, bold), ...]"""
    p = doc.add_paragraph()
    for text, size, color, bold in parts:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.bold = bold
    return p

def section_heading(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x1A, 0x27, 0x44)
    r.bold = True
    r.underline = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)

# ═══ HEADER ═══
add_title('Thasin Jahangir', 22)
add_mixed_line([
    ('Consultant QA Senior · Freelance SASU', 13, RGBColor(0x1A, 0x4A, 0x9E), True)
])
add_mixed_line([
    ('Paris & Remote', 9, RGBColor(0x66,0x66,0x66), False),
    (' · ', 9, RGBColor(0x66,0x66,0x66), False),
    ('07 69 78 72 13', 9, RGBColor(0x66,0x66,0x66), False),
    (' · ', 9, RGBColor(0x66,0x66,0x66), False),
    ('thasin@live.com', 9, RGBColor(0x66,0x66,0x66), False),
    (' · ', 9, RGBColor(0x66,0x66,0x66), False),
    ('linkedin.com/in/thasin-j-47582635', 9, RGBColor(0x66,0x66,0x66), False),
])

# ═══ PROFIL ═══
section_heading('Profil')
p = doc.add_paragraph(
    "Consultant QA senior avec 12 ans d'expérience en recette fonctionnelle. "
    "J'interviens sur des projets complexes (web app, applications mobiles Android & iOS, "
    "ERP, bancaire) pour des clients comme BRED, Accor, Oodrive, Vinci Construction, Visiodent, "
    "avec une approche Agile, collaborative et orientée résultats."
)

# ═══ COMPÉTENCES ═══
section_heading('Compétences')
skills = [
    ("Pilotage QA & Gouvernance",
     "Stratégie de test · Plans de recette · Coordination QA · UAT · Non-régression · "
     "Reporting qualité · Go/No-Go · Gestion des anomalies · Suivi d'avancement · "
     "Couverture de tests · Tableaux de bord Jira · Indicateurs qualité · Agile / Scrum"),
    ("Recette fonctionnelle",
     "Tests web & mobile · Rédaction de cas de test depuis les user stories · "
     "Exécution de campagnes de validation · Validation de parcours critiques · "
     "Recette fonctionnelle · Gherkin · Appium (bases légères) · Ranorex"),
    ("Outils & Environnements",
     "Jira · Xray · Zephyr · TestRail · Confluence · ServiceNow · SQL Oracle · "
     "Mainframe · Charles Proxy · Crashlytics · Dynatrace · Redmine · "
     "SAP BPC · SAP HANA · Windows / macOS / Linux · Android / iOS · VMware"),
    ("IA & Outils techniques",
     "Agents IA · Agentique · Prompt engineering · MCP · Claude Code · Python · Flask · "
     "GitHub Actions · LLMs cloud et locaux · Ollama · LM Studio · VS Code · Xcode"),
    ("Savoir-être",
     "Rigueur · Sens du détail · Esprit analytique · Autonomie · Proactivité · "
     "Communication claire · Collaboration avec les équipes métier et techniques · "
     "Capacité de priorisation · Curiosité technique · Orientation qualité")
]
for title, items in skills:
    p = doc.add_paragraph()
    r = p.add_run(title + '  ')
    r.bold = True
    r.font.size = Pt(9.5)
    r2 = p.add_run(items)
    r2.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(1)

# ═══ EXPÉRIENCES ═══
section_heading('Expériences')

# Sabbatique en premier (comme dans le PDF)
experiences = [
    ("Année sabbatique — montée en compétences & projets personnels", "Avr. 2025 — Présent",
     "Freelance SASU", [
        "Exploré 7 pays d'Asie du Sud-Est : Thaïlande, Malaisie, Cambodge, Indonésie, Philippines, Vietnam, Laos — immersion internationale.",
        "Monter en compétences sur IA générative, prompt engineering et digital product management (formations continues).",
        "Développer des projets agentiques, assistants autonomes et chasse d'emploi pilotée par IA.",
        "Renforcer la stack technique sur Python, Flask, GitHub Actions, MCP et LLMs cloud / locaux.",
     ], "Python, Flask, Hermes Agent, DeepSeek, GitHub Actions, MCP, Ollama, LM Studio"),
    ("BRED Banque Populaire — Facturation électronique", "Avr. 2023 — Avr. 2025",
     "Consultant QA Senior — Freelance SASU", [
        "Piloter la recette fonctionnelle sur un programme de conformité réglementaire — dématérialisation des factures, périmètres P2P/O2C.",
        "Concevoir, planifier et exécuter des campagnes de tests fonctionnels, de régression ; rédiger et maintenir les plans de test, cahiers de recette, cas de tests, scénarios de test et matrices de couverture en priorisant selon les risques et l'impact métier.",
        "Coordonner des campagnes multi-équipes dans Jira/Xray, suivi de couverture, gestion des anomalies et reporting Go/No-Go.",
        "Synchroniser avec équipes métier, BA, développeurs et support en Agile / Scrum.",
        "Sécuriser les cycles de non-régression dans un environnement Mainframe / SQL Oracle.",
     ], "Jira, Xray, SQL Oracle, Mainframe, Corcentric, ServiceNow"),
    ("Accor — Applications mobiles internationales", "Sept. 2019 — Déc. 2022",
     "Référent QA Mobile — Consultant", [
        "Référent QA sur applications Android/iOS déployées dans 100+ pays : validation fonctionnelle, UAT et non-régression.",
        "Concevoir, planifier et exécuter des campagnes de tests fonctionnels, de régression et de compatibilité multi-versions Android/iOS ; rédiger et maintenir les plans de test, cahiers de recette, cas de tests, scénarios de test en priorisant selon les risques et l'impact métier.",
        "Co-construire les user stories avec PO et développeurs pour renforcer la testabilité.",
        "Formaliser les scénarios critiques en Gherkin pour l'alignement avec l'automatisation mobile.",
        "Valider les flux transactionnels multi-appareils/multi-langues, reporting à la direction produit.",
     ], "Jira, Xray, Gherkin, Charles Proxy, Crashlytics, Android, iOS, Confluence"),
    ("Oodrive — Cloud souverain B2B", "Janv. 2017 — Avr. 2019",
     "Ingénieur Qualité Logiciel / Responsable QA", [
        "Responsable qualité sur applications critiques B2B/B2C cloud et desktop : WebSynchro (synchronisation fichiers), PostFiles (partage sécurisé), BoardNox (solution de gouvernance pour comités de direction / CA).",
        "Concevoir, planifier et exécuter des campagnes de tests fonctionnels, de régression, de compatibilité multi-navigateurs et responsive ; rédiger et maintenir les plans de test, cahiers de recette, cas de tests, scénarios de test en priorisant selon les risques et l'impact métier.",
        "Structurer les pratiques QA : templates, critères, gestion des anomalies dans Jira/Zephyr.",
        "Support N3 grands comptes, introduction automatisation Ranorex, et collaboration équipes dev/support/PO.",
     ], "Jira, Zephyr, Ranorex, Confluence, Windows, macOS, Linux"),
    ("Visiodent — Logiciels santé", "Juin — Août 2019",
     "Responsable Qualité Logiciel", [
        "Encadrer une équipe de testeurs et superviser un prestataire externe d'automatisation sur le logiciel métier santé Veasy.",
        "Concevoir, planifier et exécuter des campagnes de tests fonctionnels, de régression et de compatibilité multi-navigateurs sur les parcours métier et réglementaires.",
        "Rédiger et maintenir les plans de test, cahiers de recette, cas de tests, scénarios de test et matrices de couverture, prioriser selon les risques fonctionnels et l'impact métier.",
        "Intégrer la Carte Vitale dans les scénarios de test pour couvrir les exigences réglementaires.",
        "Coordonner les validations avec développeurs, PO et équipes métier.",
     ], "Jira, Xray, Dynatrace, CloudNetCare, Excel"),
    ("Vinci Construction — SAP BPC", "Oct. 2016 — Janv. 2017",
     "Consultant QA / Ingénieur Qualité", [
        "Recette fonctionnelle d'applications SAP BPC : conformité, intégrité des données et cohérence des états financiers.",
        "Rédiger des cas de tests depuis les besoins métier, élaborer les plans de test, et exécuter les campagnes de validation sur les applicatifs SAP BPC.",
        "Collaborer avec les équipes métier pour sécuriser les validations.",
     ], "Redmine, TestRail, SAP BPC, SAP HANA"),
    ("Profil Technology — sécurité & filtrage web", "Janv. 2011 — Sept. 2016",
     "Ingénieur Tests & Support Niveau 3", [
        "Rédiger des cas de tests depuis les users stories, élaborer les plans de test, et exécuter les campagnes de validation sur les solutions de sécurité et filtrage web.",
        "Conduire des campagnes de tests fonctionnels et non-régression multi-plateformes (Witigo, Bitdefender, WebFilter).",
        "Développer des scripts d'automatisation pour accélérer les validations et la détection de régressions.",
        "Gérer un laboratoire de tests sécurité sur Windows, macOS, Linux, VMware.",
        "Support N3 clients grands comptes en lien avec développement et support.",
     ], "VMware, Windows, macOS, Linux, Bitdefender"),
]

for title, date, role, bullets, env in experiences:
    add_mixed_line([
        (title, 10.5, RGBColor(0x11,0x11,0x11), True),
        (f"    {date}", 9, RGBColor(0x88,0x88,0x88), False),
    ])
    add_mixed_line([
        (role, 9.5, RGBColor(0x1A,0x4A,0x9E), True),
    ])
    for b in bullets:
        bp = doc.add_paragraph(b, style='List Bullet')
        bp.paragraph_format.space_after = Pt(0)
        bp.paragraph_format.space_before = Pt(0)
        for run in bp.runs:
            run.font.size = Pt(9)
    ep = doc.add_paragraph(f"Environnement : {env}")
    for run in ep.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x88,0x88,0x88)
    ep.paragraph_format.space_after = Pt(4)

# ═══ PROJETS PERSONNELS ═══
section_heading('Projets Personnels')

projects = [
    ("Pipelines QA Agentic — orchestration multi-agents",
     "Workflows agentiques pour la QA : délégation de tâches, génération de tests par IA, "
     "analyse de régression assistée. Hermes Agent, MCP, Claude Code, Codex CLI, DeepSeek, Gemma.",
     "https://nous-daily.vercel.app/"),
    ("Nous AI News — agrégation IA temps réel",
     "Pipeline ingestion 86 flux RSS avec extraction d'entités et catégorisation LLM. "
     "Moteur d'auto-amélioration continue (re-scoring, boucles RLHF, optimisation prompts auto). "
     "Next.js 14 + Supabase, 263 tests. CI/CD GitHub Actions. TypeScript, Tailwind, OpenAI GPT-5 nano.",
     "https://jobhunt-1-mt17.onrender.com/"),
    ("JobHunt — dashboard QA piloté par IA",
     "Scheduler cron autonome orchestrant scraping horaire, alertes WhatsApp, bilans hebdomadaires. "
     "Pipeline de scraping multi-sources (LinkedIn, Free-Work, Indeed, Malt, 8+ plateformes). "
     "Moteur de matching CV/offres avec scoring /100, détection de doublons et analyse TJM par marché. "
     "Auto-enrichissement IA des offres (tech stack, salaire). Génération automatisée de lettres de motivation. "
     "Alertes WhatsApp temps réel, rapports hebdomadaires. Dashboard temps réel : couverture, filtres, stats. "
     "Suite complète de tests : 83 tests unitaires/intégration, 19 scénarios BDD Gherkin, tests Playwright E2E. "
     "CI/CD GitHub Actions + Render. Auto-healing du pipeline. Python, Flask, DeepSeek.",
     "https://github.com/AtmanTest/jobhunt"),
]

for title, desc, url in projects:
    add_mixed_line([
        (title, 10, RGBColor(0x11,0x11,0x11), True),
    ])
    p = doc.add_paragraph(desc)
    for run in p.runs:
        run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(0)
    if url:
        up = doc.add_paragraph(url)
        for run in up.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x1A, 0x4A, 0x9E)
            run.underline = True
        up.paragraph_format.space_after = Pt(4)

# ═══ FORMATIONS ═══
section_heading('Formations')
formations = [
    ("ISTQB — Fondation du test logiciel", "Learning Tree", "2015"),
]
for title, school, years in formations:
    add_mixed_line([
        (title, 9.5, RGBColor(0x11,0x11,0x11), True),
        (f" — {school}", 9, RGBColor(0x66,0x66,0x66), False),
        (f" ({years})", 9, RGBColor(0x88,0x88,0x88), False),
    ])

# ═══ CERTIFICATIONS ═══
section_heading('Certifications')
certs = [
    "PSPO — Professional Scrum Product Owner · Scrum.org, 2020",
    "IA Générative & Prompt Engineering · Google / Coursera, 2025",
    "Digital Product Management · Coursera, 2025",
    "Introduction to Python · Coursera, 2025",
    "Introduction to AI · Coursera, 2025",
    "Basic SQL Syntax · Coursera, 2025",
    "Becoming an SAP Professional · Coursera, 2025",
    "Emotional Intelligence in Leadership · Coursera, 2025",
    "API Testing Journey With Postman Tool · Coursera, 2025",
]
for c in certs:
    p = doc.add_paragraph(c, style='List Bullet')
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    for run in p.runs:
        run.font.size = Pt(9)

# ═══ LANGUES ═══
section_heading('Langues')
langs = [
    ("Français", "natif"),
    ("Anglais", "courant (C1/C2)"),
    ("Espagnol", "professionnel"),
    ("Bengali", "courant"),
]
p = doc.add_paragraph()
for i, (lang, level) in enumerate(langs):
    if i > 0:
        p.add_run('  •  ').font.size = Pt(9)
    r = p.add_run(lang)
    r.font.size = Pt(9.5)
    r.bold = (lang == "Anglais")
    r2 = p.add_run(f' — {level}')
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# Save
outpath = '/Users/jahangir/jobhunt/cv_tasin_jahangir.docx'
doc.save(outpath)
print(f'✅ Word ready: {outpath}')
